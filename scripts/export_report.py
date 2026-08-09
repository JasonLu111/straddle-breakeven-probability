"""Render reports/research_report.md content into polished .docx and .pdf versions
for portfolio / print use. The Markdown file remains the source of truth for the
GitHub-hosted version; this script re-renders the same Phase 1 findings with
richer layout (title page, embedded figures, formatted tables) since Markdown
tables/figures don't carry over 1:1 into Word/PDF.
"""
import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parents[1]))

import pandas as pd

CJK_FONT_PATH = "C:/Windows/Fonts/msjh.ttc"  # Microsoft JhengHei (Traditional Chinese)
FIGURES_DIR = Path("reports/figures")
TABLES_DIR = Path("reports/tables")
OUT_DIR = Path("reports")

TITLE = "Straddle Breakeven Probability Lab"
SUBTITLE = "跨式策略損益兩平機率研究：以波動率壓縮與機率校準模型為核心"
REPORT_TITLE = "Phase 1-4 研究報告（完整版）：市場事件、Breakeven 資料集、機率模型與策略比較"


def load_summary_table() -> pd.DataFrame:
    df = pd.read_csv(TABLES_DIR / "h1_summary_table.csv")
    return df


def load_h2_summary_table() -> pd.DataFrame:
    return pd.read_csv(TABLES_DIR / "h2_summary_table.csv")


def load_phase3_table() -> pd.DataFrame:
    return pd.read_csv(TABLES_DIR / "phase3_model_comparison.csv")


def load_phase4_stats() -> pd.DataFrame:
    return pd.read_csv("results/backtests/phase4_strategy_stats.csv")


def load_phase4_pairwise() -> pd.DataFrame:
    return pd.read_csv("results/backtests/phase4_pairwise_tests.csv")


PRIMARY_ROWS_PHASE3 = [
    ("dummy_prior（無條件歷史發生率）", "ROC-AUC 0.437, Brier 0.2508"),
    ("compression_rule（單一 compression 規則）", "ROC-AUC 0.437, Brier 0.2527"),
    ("logistic_regression（22 特徵）", "ROC-AUC 0.502, Brier 0.2521"),
    ("random_forest（22 特徵，300 棵樹）", "ROC-AUC 0.516, Brier 0.2533"),
]


PRIMARY_ROWS_H2 = [
    ("有效交易週數", "681"),
    ("Compression 週數", "74"),
    ("P(breakeven | compression)", "41.9%  (95% bootstrap CI: [31.1%, 52.7%])"),
    ("P(breakeven) 無條件", "41.6%  (95% bootstrap CI: [37.7%, 45.4%])"),
    ("Probability ratio", "1.008"),
    ("平均 net PnL (compression)", "-$195.56 / 口"),
    ("平均 net PnL (normal)", "-$38.89 / 口"),
    ("Welch's t-test (PnL 差異)", "p = 0.088"),
]


PRIMARY_ROWS = [
    ("觀察數（compression / normal）", "606 / 4,536"),
    ("平均未來 20 日絕對報酬 (compression)", "2.90%"),
    ("平均未來 20 日絕對報酬 (normal)", "3.52%"),
    ("平均差異 (compression - normal)", "-0.63%  (95% bootstrap CI: [-0.81%, -0.44%])"),
    ("Welch's t-test", "t = -6.64, p = 4.85e-11"),
    ("Mann-Whitney U test", "p = 0.018"),
    ("Benjamini-Hochberg (12 規格聯合檢定)", "拒絕虛無假設 (reject)"),
    ("P(大幅波動 | compression)", "17.3%"),
    ("P(大幅波動) 無條件機率", "26.1%"),
    ("Probability ratio", "0.66"),
]


def build_docx() -> Path:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt, RGBColor

    doc = Document()

    def set_cjk(run, font_name="Microsoft JhengHei", size=None, bold=None):
        run.font.name = font_name
        run.font.element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
        if size:
            run.font.size = Pt(size)
        if bold is not None:
            run.font.bold = bold
        return run

    # Title page
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(TITLE)
    set_cjk(r, size=26, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(SUBTITLE)
    set_cjk(r, size=15)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(REPORT_TITLE)
    set_cjk(r, size=18, bold=True)
    r.font.color.rgb = RGBColor(0x40, 0x40, 0x40)

    doc.add_paragraph()

    def h(text, level=1):
        p = doc.add_heading(level=level)
        r = p.add_run(text)
        set_cjk(r, size=18 if level == 1 else 14, bold=True)

    def body(text, size=11):
        p = doc.add_paragraph()
        r = p.add_run(text)
        set_cjk(r, size=size)
        return p

    h("重要聲明 (Data Tier Disclaimer)")
    body(
        "本報告為 Phase 1（市場事件研究）成果，僅使用標的股價資料（SPY, QQQ），"
        "完全未使用選擇權資料。因此本報告不構成、也不宣稱驗證了 Long Straddle 策略的"
        "真實獲利能力，正確定位為：Breakeven-event research under estimated conditions "
        "(price-only stage), 尚未涉及 option premium。"
    )

    h("1. 資料與方法")
    for label, val in [
        ("標的", "SPY, QQQ"),
        ("資料來源", "Yahoo Finance (yfinance)，日線 OHLCV + Adjusted Close"),
        ("資料期間", "2005-01-03 至 2026-08-07（5,433 個交易日）"),
        ("資料層級", "MVP tier — 僅標的價格資料"),
    ]:
        body(f"• {label}：{val}")

    h("2. 研究假說 H1")
    body(
        "H1：當過去 20 日 realized volatility 位於歷史低分位，且 Bollinger Band width "
        "較低時，未來 10 日或 20 日的絕對報酬較高。"
    )
    body(
        "Compression regime 定義：rv_20_percentile <= threshold 且 "
        "bb_width_percentile <= threshold（主要設定 threshold = 20th percentile，"
        "10th / 30th percentile 作為穩健性檢查）。"
    )

    h("3. 主要結果 (Primary spec: SPY, 20th percentile, 20D horizon)")
    table = doc.add_table(rows=1, cols=2)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    set_cjk(hdr[0].paragraphs[0].add_run("統計量"), bold=True)
    set_cjk(hdr[1].paragraphs[0].add_run("數值"), bold=True)
    for label, val in PRIMARY_ROWS:
        row = table.add_row().cells
        set_cjk(row[0].paragraphs[0].add_run(label))
        set_cjk(row[1].paragraphs[0].add_run(val))

    doc.add_paragraph()
    body(
        "結論：H1 在本樣本中被拒絕，且方向與假說相反。波動壓縮狀態下，未來 20 日絕對"
        "報酬不僅沒有變大，反而顯著更小；壓縮狀態下出現大幅波動的機率只有無條件機率的"
        "66%，而非更高。",
        size=11,
    )

    h("4. 穩健性分析（12 規格全部同方向）")
    body(
        "對 2 個標的 x 2 個預測期間 x 3 個 compression 門檻，共 12 個規格全部執行相同"
        "檢定。全部 12 個規格中，compression regime 的平均未來絕對報酬皆低於 normal "
        "regime（差異介於 -0.30% 至 -0.89% 之間），且經 Benjamini-Hochberg 校正後全部"
        "達統計顯著。方向一致，不是單一門檻或單一標的的偶然結果。"
    )

    sub_df = load_summary_table()
    cols = ["ticker", "horizon_days", "compression_threshold_pct",
            "group_mean_difference", "group_welch_p_value",
            "prob_p_conditional", "prob_p_unconditional"]
    headers = ["Ticker", "Horizon", "Threshold", "Mean diff", "Welch p", "P(large|comp)", "P(large) uncond"]
    table2 = doc.add_table(rows=1, cols=len(headers))
    table2.style = "Light Grid Accent 1"
    for i, htext in enumerate(headers):
        set_cjk(table2.rows[0].cells[i].paragraphs[0].add_run(htext), bold=True, size=9)
    for _, rowdata in sub_df[cols].iterrows():
        cells = table2.add_row().cells
        vals = [
            str(rowdata["ticker"]), f"{int(rowdata['horizon_days'])}D",
            f"{int(rowdata['compression_threshold_pct'])}%",
            f"{rowdata['group_mean_difference']*100:.2f}%",
            f"{rowdata['group_welch_p_value']:.2e}",
            f"{rowdata['prob_p_conditional']*100:.1f}%",
            f"{rowdata['prob_p_unconditional']*100:.1f}%",
        ]
        for i, v in enumerate(vals):
            set_cjk(cells[i].paragraphs[0].add_run(v), size=9)

    doc.add_page_break()
    h("5. 圖表")
    for fname, caption in [
        ("h1_forward_return_by_regime.png", "圖 1：Compression vs normal regime 未來 20 日絕對報酬分布"),
        ("compression_regime_timeline.png", "圖 2：SPY 價格走勢與 compression regime 標記"),
        ("conditional_probability_by_threshold.png", "圖 3：三種門檻下條件機率 vs 無條件機率"),
    ]:
        doc.add_picture(str(FIGURES_DIR / fname), width=Cm(15))
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cjk(cap.add_run(caption), size=9)
        doc.add_paragraph()

    h("6. 解讀")
    body(
        "低波動狀態在統計上具有相當的持續性（volatility clustering / persistence）：波動"
        "率低的時期，後續一段時間的波動率傾向於維持偏低，而非立即反轉為大幅波動，這與"
        "GARCH 類波動率聚集現象一致，但與「壓縮後必噴出」的交易直覺相反。"
    )
    body(
        "這不代表整條研究鏈失敗：真正與策略有關的是 H2（能否突破選擇權隱含的 breakeven "
        "門檻），因為門檻本身由目前已支付的權利金（隱含波動率）決定——若壓縮狀態下 IV "
        "同步偏低，即使未來絕對報酬較小，仍可能相對於便宜的權利金更容易突破。此問題須"
        "留到 Phase 2（取得真實選擇權資料後）才能檢驗。"
    )

    doc.add_page_break()
    h("Part B: Phase 2 — Breakeven Dataset（真實選擇權資料）")

    h("1. 資料與方法", level=2)
    for label, val in [
        ("資料來源", "ORATS Data API, History (EOD) 方案"),
        ("資料層級", "Full-research tier — 真實 bid/ask, IV, delta, open interest"),
        ("回測期間", "2013-01-01 至 2026-08-07（實測驗證：2013 年前 20-40 DTE 合約覆蓋率過低）"),
        ("進場規則", "每週五，最接近 ATM strike，20-40 DTE 中最接近 30 DTE 的到期日，持有至到期"),
    ]:
        body(f"• {label}：{val}")
    body(
        "開發過程中發現並修正兩個資料正確性問題：(1) ORATS 到期日採 OCC 慣例標示為"
        "結算星期六而非實際交易日星期五，已修正為自動回溯最近交易日；(2) 到期損益計算"
        "原本誤用股息調整後價格（adj_close），造成長年股息調整累積的假性大幅跳空，已"
        "修正為使用未調整實際成交價（close）。詳見 reports/limitations.md。"
    )

    h("2. 主要結果：H2（SPY, 20th percentile threshold）", level=2)
    table3 = doc.add_table(rows=1, cols=2)
    table3.style = "Light Grid Accent 1"
    hdr3 = table3.rows[0].cells
    set_cjk(hdr3[0].paragraphs[0].add_run("統計量"), bold=True)
    set_cjk(hdr3[1].paragraphs[0].add_run("數值"), bold=True)
    for label, val in PRIMARY_ROWS_H2:
        row = table3.add_row().cells
        set_cjk(row[0].paragraphs[0].add_run(label))
        set_cjk(row[1].paragraphs[0].add_run(val))

    doc.add_paragraph()
    body(
        "結論：H2 沒有被支持，是乾淨的虛無結果。Compression 下的條件 breakeven 機率"
        "（41.9%）與無條件機率（41.6%）幾乎相同，95% bootstrap 信賴區間高度重疊。這與"
        "H1（未來絕對報酬顯著更小）形成對比：compression 狀態下 IV／權利金通常也同步"
        "偏低，breakeven 門檻跟著收窄，兩個效應大致互相抵銷。"
    )

    h("3. 穩健性分析（6 規格皆無顯著差異，但方向一致偏負）", level=2)
    sub_df2 = load_h2_summary_table()
    cols2 = ["ticker", "compression_threshold_pct", "prob_n_condition_true",
             "prob_p_conditional", "prob_p_unconditional", "prob_probability_ratio",
             "pnl_mean_treatment", "pnl_mean_control", "pnl_welch_p_value"]
    headers2 = ["Ticker", "Thresh.", "n(comp)", "P(be|comp)", "P(be) uncond", "Ratio", "PnL(comp)", "PnL(normal)", "Welch p"]
    table4 = doc.add_table(rows=1, cols=len(headers2))
    table4.style = "Light Grid Accent 1"
    for i, htext in enumerate(headers2):
        set_cjk(table4.rows[0].cells[i].paragraphs[0].add_run(htext), bold=True, size=9)
    for _, rowdata in sub_df2[cols2].iterrows():
        cells = table4.add_row().cells
        vals = [
            str(rowdata["ticker"]), f"{int(rowdata['compression_threshold_pct'])}%",
            f"{int(rowdata['prob_n_condition_true'])}",
            f"{rowdata['prob_p_conditional']*100:.1f}%",
            f"{rowdata['prob_p_unconditional']*100:.1f}%",
            f"{rowdata['prob_probability_ratio']:.2f}",
            f"${rowdata['pnl_mean_treatment']:.0f}",
            f"${rowdata['pnl_mean_control']:.0f}",
            f"{rowdata['pnl_welch_p_value']:.3f}",
        ]
        for i, v in enumerate(vals):
            set_cjk(cells[i].paragraphs[0].add_run(v), size=9)

    doc.add_paragraph()
    body(
        "有一個值得留意但未達顯著的一致方向：6 個規格中，compression 狀態下的平均 net "
        "PnL 全部比 normal 狀態更差。沒有任何一個單獨達到 5% 顯著水準且未經多重檢定"
        "校正，樣本數也偏小（compression 子樣本僅 26-145 筆），應留待 Phase 3/4 用更大"
        "樣本或更精細的模型重新檢驗，不應直接當作可交易訊號。"
    )

    doc.add_page_break()
    h("4. 圖表", level=2)
    for fname, caption in [
        ("h2_breakeven_probability.png", "圖 4：Compression vs normal regime 的真實 breakeven 機率"),
        ("h2_net_pnl_by_regime.png", "圖 5：兩種 regime 下的真實 net PnL 分布"),
    ]:
        doc.add_picture(str(FIGURES_DIR / fname), width=Cm(14))
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cjk(cap.add_run(caption), size=9)
        doc.add_paragraph()

    h("5. 下一步", level=2)
    body(
        "Phase 3 的多變量機率模型必須納入 option-cost features（iv_minus_rv、"
        "straddle_premium_pct 等），不能只用價格壓縮特徵，因為 H2 已證明單一 compression "
        "規則對真實 breakeven 機率沒有單變量預測力。Compression 下 PnL 偏負但未顯著的"
        "觀察，可作為 Phase 3 模型是否真的捕捉到交互作用的檢查點。"
    )

    doc.add_page_break()
    h("Part C: Phase 3 — Probability Model（H3-H5）")

    h("1. 資料與方法", level=2)
    for label, val in [
        ("特徵", "22 個變數：波動率、壓縮、動能、選擇權成本四大類（見 configs/model.yaml）"),
        ("標籤", "target_expiry（Phase 2 的真實 breakeven 事件）"),
        ("驗證方式", "Walk-forward expanding window：初始訓練 6 年，之後每次測試 1 年，共 8 個 fold（2019-2026 上半年）"),
        ("模型", "dummy_prior、compression_rule、logistic_regression（22 特徵）、random_forest（300 棵淺樹）"),
        ("校準", "Sigmoid (Platt) 為主，isotonic 為對照（樣本量小，isotonic 較不穩定）"),
        ("超參數", "事先固定於 configs/model.yaml，不在 fold 內調參，避免疊加資料窺探風險"),
    ]:
        body(f"• {label}：{val}")
    body(
        "所有 scaling、calibration（內部 5-fold CV）與模型訓練，都嚴格限制在該 fold 的"
        "訓練窗口內完成，測試窗口的任何一筆資料都不會被用於挑選特徵、校準或調參。"
    )

    h("2. 主要結果（SPY, pooled OOS, n=378）", level=2)
    table5 = doc.add_table(rows=1, cols=2)
    table5.style = "Light Grid Accent 1"
    hdr5 = table5.rows[0].cells
    set_cjk(hdr5[0].paragraphs[0].add_run("模型"), bold=True)
    set_cjk(hdr5[1].paragraphs[0].add_run("結果"), bold=True)
    for label, val in PRIMARY_ROWS_PHASE3:
        row = table5.add_row().cells
        set_cjk(row[0].paragraphs[0].add_run(label))
        set_cjk(row[1].paragraphs[0].add_run(val))

    doc.add_paragraph()
    body(
        "結論：四個模型的樣本外判別力都接近隨機猜測（ROC-AUC 0.50 為隨機基準），沒有"
        "一個模型展現出可靠的預測力。Random Forest 名目最高但差距極小，且 Brier score "
        "幾乎與最無資訊量的 dummy_prior 基準相同。這與 Phase 2 的 H2 結果完全一致："
        "增加模型複雜度並不能無中生有創造出資料本身不存在的訊號。"
    )

    doc.add_page_break()
    h("3. 圖表", level=2)
    for fname, caption in [
        ("phase3_calibration_curve.png", "圖 6：四個模型（SPY, sigmoid 校準）的 calibration curve"),
        ("phase3_roc_curve.png", "圖 7：四個模型的 ROC curve（pooled OOS）"),
        ("phase3_feature_importance.png", "圖 8：Logistic Regression 係數與 Random Forest feature importance"),
    ]:
        doc.add_picture(str(FIGURES_DIR / fname), width=Cm(15))
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cjk(cap.add_run(caption), size=9)
        doc.add_paragraph()

    h("4. 下一步", level=2)
    body(
        "既然 Phase 3 沒有找到可靠的機率判別力，Phase 4 的 probability-filtered "
        "strategy 預期不會顯著優於無條件買進或 compression 規則。但仍值得執行 Phase "
        "4：策略績效與機率判別力是不同的評估維度，門檻篩選仍可能透過避開少數極端虧損"
        "交易而改善風險調整後績效，需要實際跑過回測才能確認或證偽。"
    )

    doc.add_page_break()
    h("Part D: Phase 4 — Strategy Comparison（H6）")

    h("1. 方法", level=2)
    body(
        "比較三種進場策略：A（無條件買進）、B（僅在 compression regime 進場）、"
        "C（機率篩選，Phase 3 Logistic Regression sigmoid 校準預測機率 > 該 fold 訓練集"
        "自身預測機率中位數）。三者都限制在 2019-2026 共同窗口（與 Phase 3 walk-forward "
        "測試期間相同），因為 Strategy C 只有這段期間有誠實的樣本外機率。Strategy C 的"
        "門檻只用該 fold 的訓練集資料計算，套用到下一個測試窗口，未曾使用測試期資料。"
    )

    h("2. 主要結果（共同窗口，n=378）", level=2)
    phase4_stats = load_phase4_stats()
    main_strats = ["A_unconditional", "B_compression_rule", "C_probability_filtered"]
    cols4 = ["ticker", "strategy", "n_trades", "win_rate", "avg_net_pnl", "max_drawdown",
             "cvar_5pct", "longest_losing_streak"]
    headers4 = ["Ticker", "Strategy", "Trades", "Win rate", "Avg PnL", "Max DD", "CVaR(5%)", "Streak"]
    table6 = doc.add_table(rows=1, cols=len(headers4))
    table6.style = "Light Grid Accent 1"
    for i, htext in enumerate(headers4):
        set_cjk(table6.rows[0].cells[i].paragraphs[0].add_run(htext), bold=True, size=9)
    for _, rowdata in phase4_stats[phase4_stats.strategy.isin(main_strats)][cols4].iterrows():
        cells = table6.add_row().cells
        vals = [
            str(rowdata["ticker"]), str(rowdata["strategy"]).replace("_", " "),
            f"{int(rowdata['n_trades'])}", f"{rowdata['win_rate']*100:.1f}%",
            f"${rowdata['avg_net_pnl']:.0f}", f"${rowdata['max_drawdown']:.0f}",
            "n/a" if pd.isna(rowdata["cvar_5pct"]) else f"${rowdata['cvar_5pct']:.0f}",
            f"{int(rowdata['longest_losing_streak'])}",
        ]
        for i, v in enumerate(vals):
            set_cjk(cells[i].paragraphs[0].add_run(v), size=8.5)

    doc.add_paragraph()
    body(
        "結論：H6 沒有被支持。Strategy C 與 Strategy A 在統計上無法區分（SPY 與 QQQ 的"
        "成對檢定 p 值皆遠高於 0.05），與 Phase 3 的近乎隨機判別力一致。Strategy B 在 "
        "SPY 上顯著劣於 A（Welch p=0.015），但更適合厚尾分布的 Mann-Whitney 檢定未達"
        "顯著（p=0.086），證據強度中等。QQQ 的 Strategy C 視覺上優於 A（權益曲線見圖 "
        "9-10），但個別交易標準差（約 $1,600）遠大於兩者平均差距，統計上不顯著。"
    )

    doc.add_page_break()
    h("3. 權益曲線", level=2)
    for fname, caption in [
        ("phase4_spy_equity_curves.png", "圖 9：SPY 三種策略權益曲線（共同窗口）"),
        ("phase4_qqq_equity_curves.png", "圖 10：QQQ 三種策略權益曲線（共同窗口）"),
    ]:
        doc.add_picture(str(FIGURES_DIR / fname), width=Cm(15))
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cjk(cap.add_run(caption), size=9)
        doc.add_paragraph()

    h("4. 研究鏈總結", level=2)
    summary_rows = [
        ("H1：波動壓縮是否預測未來較大絕對報酬？", "Phase 1", "否——方向相反（12/12 規格顯著）"),
        ("H2：壓縮是否提高真實 breakeven 機率？", "Phase 2", "沒有證據（6/6 規格無顯著差異）"),
        ("H3-H5：多變量模型是否有可靠判別力？", "Phase 3", "否（pooled OOS ROC-AUC 皆接近 0.5）"),
        ("H6：機率篩選能否改善策略績效？", "Phase 4", "否（與無條件買進統計上無法區分）"),
    ]
    table7 = doc.add_table(rows=1, cols=3)
    table7.style = "Light Grid Accent 1"
    for i, htext in enumerate(["問題", "階段", "結論"]):
        set_cjk(table7.rows[0].cells[i].paragraphs[0].add_run(htext), bold=True)
    for q, p, c in summary_rows:
        row = table7.add_row().cells
        set_cjk(row[0].paragraphs[0].add_run(q), size=10)
        set_cjk(row[1].paragraphs[0].add_run(p), size=10)
        set_cjk(row[2].paragraphs[0].add_run(c), size=10)

    doc.add_paragraph()
    body(
        "這個結論本身就是這個專案的價值所在：它展示了一條完整、可重現、誠實面對虛無"
        "結果的量化研究流程——從市場假說、真實選擇權資料建構、嚴謹的樣本外驗證，到"
        "策略層級的統計檢定，每一步都沒有為了得到「漂亮的結果」而扭曲方法論或選擇性"
        "報告。低波動壓縮本身不是一個可以直接拿來交易 Long Straddle 的訊號，至少在本"
        "研究涵蓋的樣本、特徵與模型設定下是如此。"
    )

    out_path = OUT_DIR / "research_report.docx"
    doc.save(str(out_path))
    return out_path


def build_pdf() -> Path:
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import cm
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    from reportlab.platypus import (Image, PageBreak, Paragraph, SimpleDocTemplate,
                                     Spacer, Table, TableStyle)

    pdfmetrics.registerFont(TTFont("CJK", CJK_FONT_PATH, subfontIndex=0))
    pdfmetrics.registerFont(TTFont("CJK-Bold", CJK_FONT_PATH, subfontIndex=1))

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle("TitleCJK", fontName="CJK-Bold", fontSize=22, leading=28, alignment=1, spaceAfter=10)
    subtitle_style = ParagraphStyle("SubtitleCJK", fontName="CJK", fontSize=13, leading=20, alignment=1, spaceAfter=6)
    h1_style = ParagraphStyle("H1CJK", fontName="CJK-Bold", fontSize=15, spaceBefore=14, spaceAfter=8, textColor=colors.HexColor("#1a1a1a"))
    body_style = ParagraphStyle("BodyCJK", fontName="CJK", fontSize=10.5, leading=16, spaceAfter=8)
    caption_style = ParagraphStyle("CaptionCJK", fontName="CJK", fontSize=8.5, alignment=1, textColor=colors.grey, spaceAfter=14)

    story = []
    story.append(Paragraph(TITLE, title_style))
    story.append(Paragraph(SUBTITLE, subtitle_style))
    story.append(Paragraph(REPORT_TITLE, subtitle_style))
    story.append(Spacer(1, 0.6 * cm))

    story.append(Paragraph("重要聲明 (Data Tier Disclaimer)", h1_style))
    story.append(Paragraph(
        "本報告為 Phase 1（市場事件研究）成果，僅使用標的股價資料（SPY, QQQ），完全未"
        "使用選擇權資料。因此本報告不構成、也不宣稱驗證了 Long Straddle 策略的真實獲利"
        "能力，正確定位為 price-only 階段的 breakeven-event 研究，尚未涉及 option premium。",
        body_style))

    story.append(Paragraph("1. 資料與方法", h1_style))
    for label, val in [
        ("標的", "SPY, QQQ"),
        ("資料來源", "Yahoo Finance (yfinance)，日線 OHLCV + Adjusted Close"),
        ("資料期間", "2005-01-03 至 2026-08-07（5,433 個交易日）"),
        ("資料層級", "MVP tier — 僅標的價格資料"),
    ]:
        story.append(Paragraph(f"• {label}：{val}", body_style))

    story.append(Paragraph("2. 研究假說 H1", h1_style))
    story.append(Paragraph(
        "H1：當過去 20 日 realized volatility 位於歷史低分位，且 Bollinger Band width "
        "較低時，未來 10 日或 20 日的絕對報酬較高。", body_style))
    story.append(Paragraph(
        "Compression regime 定義：rv_20_percentile &lt;= threshold 且 bb_width_percentile "
        "&lt;= threshold（主要設定 threshold = 20th percentile，10th / 30th percentile "
        "作為穩健性檢查）。", body_style))

    story.append(Paragraph("3. 主要結果 (SPY, 20th percentile, 20D horizon)", h1_style))
    tbl_data = [["統計量", "數值"]] + [[k, v] for k, v in PRIMARY_ROWS]
    tbl_data = [[Paragraph(c, body_style) for c in row] for row in tbl_data]
    t = Table(tbl_data, colWidths=[8 * cm, 8.5 * cm])
    t.setStyle(TableStyle([
        ("FONTNAME", (0, 0), (-1, 0), "CJK-Bold"),
        ("FONTNAME", (0, 1), (-1, -1), "CJK"),
        ("FONTSIZE", (0, 0), (-1, -1), 9.5),
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#2c3e50")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#cccccc")),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#f5f5f5")]),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
    ]))
    story.append(t)
    story.append(Spacer(1, 0.4 * cm))
    story.append(Paragraph(
        "結論：H1 在本樣本中被拒絕，且方向與假說相反。波動壓縮狀態下，未來 20 日絕對"
        "報酬不僅沒有變大，反而顯著更小；壓縮狀態下出現大幅波動的機率只有無條件機率的"
        "66%，而非更高。", body_style))

    story.append(Paragraph("4. 穩健性分析（12 規格全部同方向）", h1_style))
    story.append(Paragraph(
        "對 2 個標的 x 2 個預測期間 x 3 個 compression 門檻，共 12 個規格全部執行相同"
        "檢定，全部方向一致並經 Benjamini-Hochberg 校正後達統計顯著。", body_style))

    sub_df = load_summary_table()
    cols = ["ticker", "horizon_days", "compression_threshold_pct",
            "group_mean_difference", "group_welch_p_value",
            "prob_p_conditional", "prob_p_unconditional"]
    headers = ["Ticker", "Horizon", "Thresh.", "Mean diff", "Welch p", "P(large|comp)", "P(large) uncond"]
    rows = [headers]
    for _, rowdata in sub_df[cols].iterrows():
        rows.append([
            str(rowdata["ticker"]), f"{int(rowdata['horizon_days'])}D",
            f"{int(rowdata['compression_threshold_pct'])}%",
            f"{rowdata['group_mean_difference']*100:.2f}%",
            f"{rowdata['group_welch_p_value']:.1e}",
            f"{rowdata['prob_p_conditional']*100:.1f}%",
            f"{rowdata['prob_p_unconditional']*100:.1f}%",
        ])
    t2 = Table(rows, colWidths=[1.8 * cm, 1.8 * cm, 1.8 * cm, 2.2 * cm, 2.4 * cm, 2.9 * cm, 2.9 * cm])
    t2.setStyle(TableStyle([
        ("FONTNAME", (0, 0), (-1, 0), "CJK-Bold"),
        ("FONTNAME", (0, 1), (-1, -1), "CJK"),
        ("FONTSIZE", (0, 0), (-1, -1), 8),
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#2c3e50")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#cccccc")),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#f5f5f5")]),
        ("ALIGN", (1, 0), (-1, -1), "CENTER"),
    ]))
    story.append(t2)
    story.append(PageBreak())

    story.append(Paragraph("5. 圖表", h1_style))
    for fname, caption in [
        ("h1_forward_return_by_regime.png", "圖 1：Compression vs normal regime 未來 20 日絕對報酬分布"),
        ("compression_regime_timeline.png", "圖 2：SPY 價格走勢與 compression regime 標記"),
        ("conditional_probability_by_threshold.png", "圖 3：三種門檻下條件機率 vs 無條件機率"),
    ]:
        img = Image(str(FIGURES_DIR / fname), width=15 * cm, height=15 * cm * 0.55)
        story.append(img)
        story.append(Paragraph(caption, caption_style))

    story.append(Paragraph("6. 解讀", h1_style))
    story.append(Paragraph(
        "低波動狀態在統計上具有相當的持續性（volatility clustering / persistence）：波動"
        "率低的時期，後續一段時間的波動率傾向於維持偏低，而非立即反轉為大幅波動，這與"
        "GARCH 類波動率聚集現象一致，但與「壓縮後必噴出」的交易直覺相反。", body_style))
    story.append(Paragraph(
        "這不代表整條研究鏈失敗：真正與策略有關的是 H2（能否突破選擇權隱含的 breakeven "
        "門檻），因為門檻本身由目前已支付的權利金（隱含波動率）決定——此問題須留到 "
        "Phase 2（取得真實選擇權資料後）才能檢驗。", body_style))

    story.append(PageBreak())
    story.append(Paragraph("Part B: Phase 2 — Breakeven Dataset（真實選擇權資料）", h1_style))

    story.append(Paragraph("1. 資料與方法", h1_style))
    for label, val in [
        ("資料來源", "ORATS Data API, History (EOD) 方案"),
        ("資料層級", "Full-research tier — 真實 bid/ask, IV, delta, open interest"),
        ("回測期間", "2013-01-01 至 2026-08-07（實測驗證：2013 年前 20-40 DTE 合約覆蓋率過低）"),
        ("進場規則", "每週五，最接近 ATM strike，20-40 DTE 中最接近 30 DTE 的到期日，持有至到期"),
    ]:
        story.append(Paragraph(f"• {label}：{val}", body_style))
    story.append(Paragraph(
        "開發過程中發現並修正兩個資料正確性問題：(1) ORATS 到期日採 OCC 慣例標示為結算"
        "星期六而非實際交易日星期五，已修正為自動回溯最近交易日；(2) 到期損益計算原本"
        "誤用股息調整後價格（adj_close），造成長年股息調整累積的假性大幅跳空，已修正為"
        "使用未調整實際成交價（close）。詳見 reports/limitations.md。", body_style))

    story.append(Paragraph("2. 主要結果：H2（SPY, 20th percentile threshold）", h1_style))
    tbl_data2 = [["統計量", "數值"]] + [[k, v] for k, v in PRIMARY_ROWS_H2]
    tbl_data2 = [[Paragraph(c, body_style) for c in row] for row in tbl_data2]
    t3 = Table(tbl_data2, colWidths=[8 * cm, 8.5 * cm])
    t3.setStyle(TableStyle([
        ("FONTNAME", (0, 0), (-1, 0), "CJK-Bold"),
        ("FONTNAME", (0, 1), (-1, -1), "CJK"),
        ("FONTSIZE", (0, 0), (-1, -1), 9.5),
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#2c3e50")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#cccccc")),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#f5f5f5")]),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
    ]))
    story.append(t3)
    story.append(Spacer(1, 0.4 * cm))
    story.append(Paragraph(
        "結論：H2 沒有被支持，是乾淨的虛無結果。Compression 下的條件 breakeven 機率"
        "（41.9%）與無條件機率（41.6%）幾乎相同，95% bootstrap 信賴區間高度重疊。這與"
        "H1（未來絕對報酬顯著更小）形成對比：compression 狀態下 IV／權利金通常也同步"
        "偏低，breakeven 門檻跟著收窄，兩個效應大致互相抵銷。", body_style))

    story.append(Paragraph("3. 穩健性分析（6 規格皆無顯著差異，但方向一致偏負）", h1_style))
    sub_df2 = load_h2_summary_table()
    cols2 = ["ticker", "compression_threshold_pct", "prob_n_condition_true",
             "prob_p_conditional", "prob_p_unconditional", "prob_probability_ratio",
             "pnl_mean_treatment", "pnl_mean_control", "pnl_welch_p_value"]
    headers2 = ["Ticker", "Thresh.", "n(comp)", "P(be|comp)", "P(be) uncond", "Ratio", "PnL(comp)", "PnL(normal)", "Welch p"]
    rows2 = [headers2]
    for _, rowdata in sub_df2[cols2].iterrows():
        rows2.append([
            str(rowdata["ticker"]), f"{int(rowdata['compression_threshold_pct'])}%",
            f"{int(rowdata['prob_n_condition_true'])}",
            f"{rowdata['prob_p_conditional']*100:.1f}%",
            f"{rowdata['prob_p_unconditional']*100:.1f}%",
            f"{rowdata['prob_probability_ratio']:.2f}",
            f"${rowdata['pnl_mean_treatment']:.0f}",
            f"${rowdata['pnl_mean_control']:.0f}",
            f"{rowdata['pnl_welch_p_value']:.3f}",
        ])
    t4 = Table(rows2, colWidths=[1.6 * cm, 1.6 * cm, 1.6 * cm, 2.0 * cm, 2.1 * cm, 1.5 * cm, 2.0 * cm, 2.0 * cm, 1.6 * cm])
    t4.setStyle(TableStyle([
        ("FONTNAME", (0, 0), (-1, 0), "CJK-Bold"),
        ("FONTNAME", (0, 1), (-1, -1), "CJK"),
        ("FONTSIZE", (0, 0), (-1, -1), 7.5),
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#2c3e50")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#cccccc")),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#f5f5f5")]),
        ("ALIGN", (1, 0), (-1, -1), "CENTER"),
    ]))
    story.append(t4)
    story.append(Spacer(1, 0.3 * cm))
    story.append(Paragraph(
        "有一個值得留意但未達顯著的一致方向：6 個規格中，compression 狀態下的平均 net "
        "PnL 全部比 normal 狀態更差。沒有任何一個單獨達到 5% 顯著水準且未經多重檢定校正，"
        "樣本數也偏小，應留待 Phase 3/4 用更大樣本重新檢驗，不應直接當作可交易訊號。",
        body_style))

    story.append(PageBreak())
    story.append(Paragraph("4. 圖表", h1_style))
    for fname, caption in [
        ("h2_breakeven_probability.png", "圖 4：Compression vs normal regime 的真實 breakeven 機率"),
        ("h2_net_pnl_by_regime.png", "圖 5：兩種 regime 下的真實 net PnL 分布"),
    ]:
        img = Image(str(FIGURES_DIR / fname), width=13 * cm, height=13 * cm * 0.65)
        story.append(img)
        story.append(Paragraph(caption, caption_style))

    story.append(Paragraph("5. 下一步", h1_style))
    story.append(Paragraph(
        "Phase 3 的多變量機率模型必須納入 option-cost features（iv_minus_rv、"
        "straddle_premium_pct 等），不能只用價格壓縮特徵，因為 H2 已證明單一 compression "
        "規則對真實 breakeven 機率沒有單變量預測力。Compression 下 PnL 偏負但未顯著的"
        "觀察，可作為 Phase 3 模型是否真的捕捉到交互作用的檢查點。", body_style))

    story.append(PageBreak())
    story.append(Paragraph("Part C: Phase 3 — Probability Model（H3-H5）", h1_style))

    story.append(Paragraph("1. 資料與方法", h1_style))
    for label, val in [
        ("特徵", "22 個變數：波動率、壓縮、動能、選擇權成本四大類（見 configs/model.yaml）"),
        ("標籤", "target_expiry（Phase 2 的真實 breakeven 事件）"),
        ("驗證方式", "Walk-forward expanding window：初始訓練 6 年，之後每次測試 1 年，共 8 個 fold（2019-2026 上半年）"),
        ("模型", "dummy_prior、compression_rule、logistic_regression（22 特徵）、random_forest（300 棵淺樹）"),
        ("校準", "Sigmoid (Platt) 為主，isotonic 為對照（樣本量小，isotonic 較不穩定）"),
        ("超參數", "事先固定於 configs/model.yaml，不在 fold 內調參，避免疊加資料窺探風險"),
    ]:
        story.append(Paragraph(f"• {label}：{val}", body_style))
    story.append(Paragraph(
        "所有 scaling、calibration（內部 5-fold CV）與模型訓練，都嚴格限制在該 fold 的"
        "訓練窗口內完成，測試窗口的任何一筆資料都不會被用於挑選特徵、校準或調參。", body_style))

    story.append(Paragraph("2. 主要結果（SPY, pooled OOS, n=378）", h1_style))
    tbl_data3 = [["模型", "結果"]] + [[k, v] for k, v in PRIMARY_ROWS_PHASE3]
    tbl_data3 = [[Paragraph(c, body_style) for c in row] for row in tbl_data3]
    t5 = Table(tbl_data3, colWidths=[8 * cm, 8.5 * cm])
    t5.setStyle(TableStyle([
        ("FONTNAME", (0, 0), (-1, 0), "CJK-Bold"),
        ("FONTNAME", (0, 1), (-1, -1), "CJK"),
        ("FONTSIZE", (0, 0), (-1, -1), 9.5),
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#2c3e50")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#cccccc")),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#f5f5f5")]),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
    ]))
    story.append(t5)
    story.append(Spacer(1, 0.4 * cm))
    story.append(Paragraph(
        "結論：四個模型的樣本外判別力都接近隨機猜測（ROC-AUC 0.50 為隨機基準），沒有"
        "一個模型展現出可靠的預測力。Random Forest 名目最高但差距極小，且 Brier score "
        "幾乎與最無資訊量的 dummy_prior 基準相同。這與 Phase 2 的 H2 結果完全一致："
        "增加模型複雜度並不能無中生有創造出資料本身不存在的訊號。", body_style))

    story.append(PageBreak())
    story.append(Paragraph("3. 圖表", h1_style))
    for fname, caption in [
        ("phase3_calibration_curve.png", "圖 6：四個模型（SPY, sigmoid 校準）的 calibration curve"),
        ("phase3_roc_curve.png", "圖 7：四個模型的 ROC curve（pooled OOS）"),
        ("phase3_feature_importance.png", "圖 8：Logistic Regression 係數與 Random Forest feature importance"),
    ]:
        img2 = Image(str(FIGURES_DIR / fname), width=14 * cm, height=14 * cm * 0.6)
        story.append(img2)
        story.append(Paragraph(caption, caption_style))

    story.append(Paragraph("4. 下一步", h1_style))
    story.append(Paragraph(
        "既然 Phase 3 沒有找到可靠的機率判別力，Phase 4 的 probability-filtered "
        "strategy 預期不會顯著優於無條件買進或 compression 規則。但仍值得執行 Phase "
        "4：策略績效與機率判別力是不同的評估維度，門檻篩選仍可能透過避開少數極端虧損"
        "交易而改善風險調整後績效，需要實際跑過回測才能確認或證偽。", body_style))

    story.append(PageBreak())
    story.append(Paragraph("Part D: Phase 4 — Strategy Comparison（H6）", h1_style))

    story.append(Paragraph("1. 方法", h1_style))
    story.append(Paragraph(
        "比較三種進場策略：A（無條件買進）、B（僅在 compression regime 進場）、C（機率"
        "篩選，Phase 3 Logistic Regression sigmoid 校準預測機率 &gt; 該 fold 訓練集自身"
        "預測機率中位數）。三者都限制在 2019-2026 共同窗口，因為 Strategy C 只有這段"
        "期間有誠實的樣本外機率。Strategy C 的門檻只用該 fold 訓練集資料計算，套用到下"
        "一個測試窗口，未曾使用測試期資料。", body_style))

    story.append(Paragraph("2. 主要結果（共同窗口，n=378）", h1_style))
    phase4_stats = load_phase4_stats()
    main_strats = ["A_unconditional", "B_compression_rule", "C_probability_filtered"]
    cols4 = ["ticker", "strategy", "n_trades", "win_rate", "avg_net_pnl", "max_drawdown",
             "cvar_5pct", "longest_losing_streak"]
    headers4 = ["Ticker", "Strategy", "Trades", "Win rate", "Avg PnL", "Max DD", "CVaR(5%)", "Streak"]
    rows4 = [headers4]
    for _, rowdata in phase4_stats[phase4_stats.strategy.isin(main_strats)][cols4].iterrows():
        rows4.append([
            str(rowdata["ticker"]), str(rowdata["strategy"]).replace("_", " "),
            f"{int(rowdata['n_trades'])}", f"{rowdata['win_rate']*100:.1f}%",
            f"${rowdata['avg_net_pnl']:.0f}", f"${rowdata['max_drawdown']:.0f}",
            "n/a" if pd.isna(rowdata["cvar_5pct"]) else f"${rowdata['cvar_5pct']:.0f}",
            f"{int(rowdata['longest_losing_streak'])}",
        ])
    t6 = Table(rows4, colWidths=[1.4 * cm, 3.3 * cm, 1.6 * cm, 1.8 * cm, 1.8 * cm, 2.0 * cm, 1.8 * cm, 1.6 * cm])
    t6.setStyle(TableStyle([
        ("FONTNAME", (0, 0), (-1, 0), "CJK-Bold"),
        ("FONTNAME", (0, 1), (-1, -1), "CJK"),
        ("FONTSIZE", (0, 0), (-1, -1), 8),
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#2c3e50")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#cccccc")),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#f5f5f5")]),
        ("ALIGN", (2, 0), (-1, -1), "CENTER"),
    ]))
    story.append(t6)
    story.append(Spacer(1, 0.4 * cm))
    story.append(Paragraph(
        "結論：H6 沒有被支持。Strategy C 與 Strategy A 在統計上無法區分（SPY 與 QQQ 的"
        "成對檢定 p 值皆遠高於 0.05），與 Phase 3 的近乎隨機判別力一致。Strategy B 在 "
        "SPY 上顯著劣於 A（Welch p=0.015），但更適合厚尾分布的 Mann-Whitney 檢定未達"
        "顯著（p=0.086），證據強度中等。QQQ 的 Strategy C 視覺上優於 A，但個別交易標準"
        "差（約 $1,600）遠大於兩者平均差距，統計上不顯著。", body_style))

    story.append(PageBreak())
    story.append(Paragraph("3. 權益曲線", h1_style))
    for fname, caption in [
        ("phase4_spy_equity_curves.png", "圖 9：SPY 三種策略權益曲線（共同窗口）"),
        ("phase4_qqq_equity_curves.png", "圖 10：QQQ 三種策略權益曲線（共同窗口）"),
    ]:
        img3 = Image(str(FIGURES_DIR / fname), width=14 * cm, height=14 * cm * 0.62)
        story.append(img3)
        story.append(Paragraph(caption, caption_style))

    story.append(Paragraph("4. 研究鏈總結", h1_style))
    summary_rows = [
        ("H1：波動壓縮是否預測未來較大絕對報酬？", "Phase 1", "否——方向相反（12/12 規格顯著）"),
        ("H2：壓縮是否提高真實 breakeven 機率？", "Phase 2", "沒有證據（6/6 規格無顯著差異）"),
        ("H3-H5：多變量模型是否有可靠判別力？", "Phase 3", "否（pooled OOS ROC-AUC 皆接近 0.5）"),
        ("H6：機率篩選能否改善策略績效？", "Phase 4", "否（與無條件買進統計上無法區分）"),
    ]
    tbl_data4 = [["問題", "階段", "結論"]] + list(summary_rows)
    tbl_data4 = [[Paragraph(c, body_style) for c in row] for row in tbl_data4]
    t7 = Table(tbl_data4, colWidths=[7 * cm, 2.5 * cm, 7 * cm])
    t7.setStyle(TableStyle([
        ("FONTNAME", (0, 0), (-1, 0), "CJK-Bold"),
        ("FONTNAME", (0, 1), (-1, -1), "CJK"),
        ("FONTSIZE", (0, 0), (-1, -1), 9),
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#2c3e50")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#cccccc")),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#f5f5f5")]),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
    ]))
    story.append(t7)
    story.append(Spacer(1, 0.4 * cm))
    story.append(Paragraph(
        "這個結論本身就是這個專案的價值所在：它展示了一條完整、可重現、誠實面對虛無"
        "結果的量化研究流程——從市場假說、真實選擇權資料建構、嚴謹的樣本外驗證，到"
        "策略層級的統計檢定，每一步都沒有為了得到「漂亮的結果」而扭曲方法論或選擇性"
        "報告。低波動壓縮本身不是一個可以直接拿來交易 Long Straddle 的訊號，至少在本"
        "研究涵蓋的樣本、特徵與模型設定下是如此。", body_style))

    out_path = OUT_DIR / "research_report.pdf"
    doc = SimpleDocTemplate(str(out_path), pagesize=A4,
                             topMargin=1.5 * cm, bottomMargin=1.5 * cm,
                             leftMargin=1.8 * cm, rightMargin=1.8 * cm)
    doc.build(story)
    return out_path


if __name__ == "__main__":
    docx_path = build_docx()
    print(f"Wrote {docx_path}")
    pdf_path = build_pdf()
    print(f"Wrote {pdf_path}")
